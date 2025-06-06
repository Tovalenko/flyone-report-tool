import streamlit as st
import pandas as pd
from deep_translator import GoogleTranslator
from datetime import datetime
from docx import Document
import io

st.set_page_config(page_title="Flyone Report Tool", layout="wide")
translator = GoogleTranslator(source='auto', target='hy')

st.title("🛫 Flyone Report Tool (Web Version)")

uploaded_excel = st.file_uploader("📁 Upload Excel File", type=["xlsx"])
start_date = st.date_input("🗕 Start Date")
end_date = st.date_input("🗕 End Date")

translations = []

if uploaded_excel:
    df = pd.read_excel(uploaded_excel, sheet_name=None)
    sheet_names = list(df.keys())
    selected_sheet = st.selectbox("📁 Choose Excel sheet", sheet_names)
    data = df[selected_sheet]

    data.columns = data.columns.str.strip()

    if 'Date & Time of Event (UTC)' in data.columns:
        data['Date & Time of Event (UTC)'] = pd.to_datetime(data['Date & Time of Event (UTC)'], errors='coerce')
        filtered = data[
            (data['Date & Time of Event (UTC)'] >= pd.to_datetime(start_date)) &
            (data['Date & Time of Event (UTC)'] <= pd.to_datetime(end_date))
        ]

        if not filtered.empty:
            st.success(f"✅ Found {len(filtered)} reports between selected dates.")

            if {'Aircraft Registration', 'Type of report'}.issubset(filtered.columns):
                grouped = filtered.groupby(['Type of report', 'Aircraft Registration'])

                for (report_type, aircraft), group in grouped:
                    similar_count = len(group)
                    st.markdown(f"### 📂 {report_type} — ✈️ {aircraft}")
                    for idx, row in group.iterrows():
                        original = str(row.get("Details", ""))
                        st.markdown(f"**📄 Original:** {original}")

                        try:
                            translated = translator.translate(original)
                            summarized = translated.strip()
                        except Exception:
                            summarized = "[Թարգմանությունն ձախողվեց]"

                        new_text = st.text_area(f"✏️ Edit Translation [{idx}]", summarized, key=f"edit_{idx}")
                        translations.append({
                            "Aircraft": aircraft,
                            "Type": report_type,
                            "Date": row.get("Date & Time of Event (UTC)"),
                            "Flight Number": row.get("Flight Number"),
                            "Departure": row.get("Departure"),
                            "Destination": row.get("Destination"),
                            "Report ID": row.get("Report ID"),
                            "Translation": new_text,
                            "Status": row.get("Status"),
                            "Similar Count": similar_count
                        })

                if st.button("📅 Export Translated Reports to Excel"):
                    export_df = pd.DataFrame(translations)
                    export_df.to_excel("translated_reports.xlsx", index=False)
                    with open("translated_reports.xlsx", "rb") as file:
                        st.download_button("⬇️ Download Excel", file, file_name="translated_reports.xlsx")
            else:
                st.error("❌ Required columns missing: 'Aircraft Registration' or 'Type of report'")
        else:
            st.warning("No reports found in selected date range.")
    else:
        st.error("❌ 'Date & Time of Event (UTC)' column not found.")

def generate_word_from_scratch(translations, start_date, end_date):
    doc = Document()

    section_titles = {
        "Ground Handling": "Վերգետնյա սպասարկում/Նստեցման հետ կապված խնտիրներ՝",
        "Technical": "Կեղկնիկական զեկույթներ՝",
        "Catering": "Քեյթերինգ",
        "Other": "Այլ զեկույթներ",
        "Cleaning": "Օդանավի աղտոտվածության վերաբերու զեկույթներ՝"
    }

    grouped = {}
    for entry in translations:
        report_type = entry["Type"]
        aircraft = entry["Aircraft"]
        grouped.setdefault(report_type, {}).setdefault(aircraft, []).append(entry)

    for report_type_en, header in section_titles.items():
        total_count = sum(len(v) for v in grouped.get(report_type_en, {}).values())
        doc.add_paragraph(f"{header} - {total_count}")

        aircraft_groups = grouped.get(report_type_en, {})
        for aircraft, entries in aircraft_groups.items():
            doc.add_paragraph(f"✈️ {aircraft}")
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Օդանավ'
            hdr_cells[1].text = 'Զեկույց N'
            hdr_cells[2].text = 'Զեկույց / Report'
            hdr_cells[3].text = 'Ամսայթ / Date'
            hdr_cells[4].text = 'Ուղղույն / Չռիչքի համար / Destination / Flight Number'
            hdr_cells[5].text = 'Զեկույցների քանակ'
            hdr_cells[6].text = 'Կարգավիփակ'

            for entry in entries:
                row_cells = table.add_row().cells
                row_cells[0].text = str(entry.get("Aircraft", ""))

                report_id = str(entry.get("Report ID", "")).strip() if report_type_en == "Technical" else ""
                row_cells[1].text = report_id

                row_cells[2].text = entry.get("Translation", "").strip()

                date = entry.get("Date")
                row_cells[3].text = date.strftime("%Y-%m-%d %H:%M") if pd.notna(date) else ""

                fn = str(entry.get("Flight Number", "") or "")
                dep = str(entry.get("Departure", "") or "")
                dest = str(entry.get("Destination", "") or "")
                route = f"{fn} / {dep}-{dest}" if dep and dest else fn
                row_cells[4].text = route.strip()

                row_cells[5].text = str(entry.get("Similar Count", ""))
                row_cells[6].text = str(entry.get("Status", ""))

        doc.add_paragraph("\n")

    output = io.BytesIO()
    filename = f"Translated_Report_{start_date.strftime('%d.%m.%y')}-{end_date.strftime('%d.%m.%y')}.docx"
    doc.save(output)
    output.seek(0)
    return output, filename

if translations and st.button("📝 Generate Word Report"):
    word_file, word_name = generate_word_from_scratch(translations, start_date, end_date)
    st.download_button("⬇️ Download Word File", word_file, file_name=word_name)
    st.success("✅ Word document created successfully.")
